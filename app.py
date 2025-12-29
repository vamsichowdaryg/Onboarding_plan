# # # from flask import Flask, request, jsonify
# # # from flask_cors import CORS
# # # import json
# # # import re
# # # from typing import List, Dict, Any
# # # import os
# # # from docx import Document
# # # import logging

# # # # Configure logging
# # # logging.basicConfig(level=logging.INFO)
# # # logger = logging.getLogger(__name__)

# # # app = Flask(__name__)
# # # CORS(app)  # Enable CORS for all routes

# # # class OnboardingCardGenerator:
# # #     """
# # #     Generator for Microsoft Teams Adaptive Cards matching the exact template structure
# # #     """
    
# # #     def __init__(self):
# # #         self.card_template = {
# # #             "$schema": "http://adaptivecards.io/schemas/adaptive-card.json",
# # #             "type": "AdaptiveCard",
# # #             "version": "1.4",
# # #             "body": [],
# # #             "actions": []
# # #         }
    
# # #     def read_docx_file(self, filepath: str) -> str:
# # #         """
# # #         Read content from DOCX file
# # #         """
# # #         try:
# # #             if not os.path.exists(filepath):
# # #                 raise FileNotFoundError(f"File not found: {filepath}")
            
# # #             doc = Document(filepath)
# # #             full_text = []
            
# # #             # Extract text from paragraphs
# # #             for para in doc.paragraphs:
# # #                 full_text.append(para.text)
            
# # #             # Extract text from tables
# # #             for table in doc.tables:
# # #                 for row in table.rows:
# # #                     row_text = []
# # #                     for cell in row.cells:
# # #                         row_text.append(cell.text)
# # #                     full_text.append(' | '.join(row_text))
            
# # #             content = '\n'.join(full_text)
# # #             logger.info(f"Successfully read document: {filepath}")
# # #             return content
            
# # #         except Exception as e:
# # #             logger.error(f"Error reading DOCX file: {str(e)}")
# # #             raise
    
# # #     def extract_questions_from_document(self, doc_content: str) -> Dict[str, List[str]]:
# # #         """
# # #         Extract questions from the actual document content
# # #         Parse the table structure and extract items
# # #         """
# # #         sections = {
# # #             "general_info": [],
# # #             "trainings": [],
# # #             "additional_items": []
# # #         }
        
# # #         lines = doc_content.split('\n')
# # #         current_section = None
# # #         in_mandatory_training = False
# # #         in_general_section = False
# # #         in_conversations = False
        
# # #         for line in lines:
# # #             line = line.strip()
            
# # #             # Detect sections
# # #             if 'General information' in line or 'to-do' in line:
# # #                 current_section = 'general'
# # #                 in_general_section = True
# # #                 in_mandatory_training = False
# # #                 in_conversations = False
# # #                 continue
# # #             elif 'Mandatory Training' in line:
# # #                 current_section = 'trainings'
# # #                 in_mandatory_training = True
# # #                 in_general_section = False
# # #                 in_conversations = False
# # #                 continue
# # #             elif 'Individual' in line and 'Conversation' in line:
# # #                 current_section = 'conversations'
# # #                 in_mandatory_training = False
# # #                 in_general_section = False
# # #                 in_conversations = True
# # #                 continue
            
# # #             # Extract general information items
# # #             if in_general_section and '|' in line:
# # #                 parts = [p.strip() for p in line.split('|') if p.strip()]
# # #                 if parts and len(parts[0]) > 20 and not any(x in parts[0].lower() for x in ['content', 'description', 'responsible', 'duration', 'due date']):
# # #                     item = parts[0].replace('**', '').strip()
# # #                     if item and len(item) > 10 and item not in sections["general_info"]:
# # #                         sections["general_info"].append(item)
            
# # #             # Extract ALL training items with checkboxes
# # #             if in_mandatory_training and '☐' in line:
# # #                 match = re.search(r'☐\s*(.+?)(?:\s*\||$)', line)
# # #                 if match:
# # #                     training_name = match.group(1).strip()
# # #                     training_name = re.sub(r'\[|\]|\{\.mark\}', '', training_name).strip()
# # #                     if training_name and len(training_name) > 2 and training_name not in sections["trainings"]:
# # #                         sections["trainings"].append(training_name)
            
# # #             # Extract conversation/additional items
# # #             if in_conversations and '☐' in line:
# # #                 match = re.search(r'☐\s*(.+?)(?:\s*\||$)', line)
# # #                 if match:
# # #                     item_name = match.group(1).strip()
# # #                     item_name = re.sub(r'\[|\]|\{\.mark\}|\*\*', '', item_name).strip()
# # #                     if item_name and len(item_name) > 5 and item_name not in sections["additional_items"]:
# # #                         sections["additional_items"].append(item_name)
        
# # #         # Use defaults matching your document if nothing extracted
# # #         if not sections["general_info"]:
# # #             sections["general_info"] = [
# # #                 "Have you attended the welcome event and received the welcome box?",
# # #                 "Was there an introduction to your manager, colleagues, and buddy?",
# # #                 "Is the workplace fully equipped and ready for work?"
# # #             ]
        
# # #         if not sections["trainings"]:
# # #             sections["trainings"] = [
# # #                 "Check-in @tesa",
# # #                 "Anti-Corruption",
# # #                 "Data-Protection",
# # #                 "Cybersecurity",
# # #                 "Code of Conduct",
# # #                 "Safety and Health at Work",
# # #                 "Sustainability Program@tesa",
# # #                 "PME @tesa",
# # #                 "tSN „to make The production of tesa tape",
# # #                 "Other function related trainings"
# # #             ]
        
# # #         if not sections["additional_items"]:
# # #             sections["additional_items"] = [
# # #                 "Meeting with department stakeholders",
# # #                 "Tour of facilities (offices, canteen, meeting areas)",
# # #                 "Introduction to current projects and priorities"
# # #             ]
        
# # #         logger.info(f"Extracted sections: {len(sections['general_info'])} general, {len(sections['trainings'])} trainings, {len(sections['additional_items'])} additional")
# # #         return sections
    
# # #     def generate_adaptive_card(self, sections: Dict[str, List[str]]) -> Dict[str, Any]:
# # #         """
# # #         Generate the complete adaptive card JSON matching the exact template
# # #         """
# # #         card = {
# # #             "$schema": "http://adaptivecards.io/schemas/adaptive-card.json",
# # #             "type": "AdaptiveCard",
# # #             "version": "1.4",
# # #             "body": []
# # #         }
        
# # #         # Container with standard width
# # #         container = {
# # #             "type": "Container",
# # #             "style": "default",
# # #             "items": []
# # #         }
        
# # #         # Header Section
# # #         header = {
# # #             "type": "ColumnSet",
# # #             "columns": [
# # #                 {
# # #                     "type": "Column",
# # #                     "width": "auto",
# # #                     "items": [
# # #                         {
# # #                             "type": "Image",
# # #                             "url": "https://via.placeholder.com/50",
# # #                             "size": "Small",
# # #                             "style": "Person",
# # #                             "altText": "NOA"
# # #                         }
# # #                     ],
# # #                     "verticalContentAlignment": "Center"
# # #                 },
# # #                 {
# # #                     "type": "Column",
# # #                     "width": "stretch",
# # #                     "items": [
# # #                         {
# # #                             "type": "TextBlock",
# # #                             "text": "NOA",
# # #                             "weight": "Bolder",
# # #                             "size": "Medium",
# # #                             "wrap": True
# # #                         },
# # #                         {
# # #                             "type": "TextBlock",
# # #                             "text": "Hiring Manager – Onboarding Template",
# # #                             "isSubtle": True,
# # #                             "spacing": "None",
# # #                             "size": "Small",
# # #                             "wrap": True
# # #                         }
# # #                     ],
# # #                     "verticalContentAlignment": "Center"
# # #                 }
# # #             ],
# # #             "separator": True,
# # #             "spacing": "Medium"
# # #         }
# # #         container["items"].append(header)
        
# # #         # General Information Section
# # #         container["items"].append({
# # #             "type": "TextBlock",
# # #             "text": "General information / to-do's",
# # #             "weight": "Bolder",
# # #             "size": "Medium",
# # #             "spacing": "Large",
# # #             "wrap": True
# # #         })
        
# # #         # General checkboxes
# # #         for idx, question in enumerate(sections["general_info"]):
# # #             container["items"].append({
# # #                 "type": "Input.Toggle",
# # #                 "id": f"general_{idx}",
# # #                 "title": question,
# # #                 "value": "false",
# # #                 "wrap": True,
# # #                 "spacing": "Small"
# # #             })
        
# # #         # Trainings Section with nested checkboxes
# # #         container["items"].append({
# # #             "type": "Container",
# # #             "items": [
# # #                 {
# # #                     "type": "Input.Toggle",
# # #                     "id": "trainings_section",
# # #                     "title": "Trainings",
# # #                     "value": "false",
# # #                     "wrap": True
# # #                 }
# # #             ],
# # #             "spacing": "Medium"
# # #         })
        
# # #         # Training items with remove buttons
# # #         for idx, training in enumerate(sections["trainings"]):
# # #             training_row = {
# # #                 "type": "ColumnSet",
# # #                 "columns": [
# # #                     {
# # #                         "type": "Column",
# # #                         "width": "auto",
# # #                         "items": [
# # #                             {
# # #                                 "type": "Input.Toggle",
# # #                                 "id": f"training_{idx}",
# # #                                 "title": "",
# # #                                 "value": "false"
# # #                             }
# # #                         ],
# # #                         "verticalContentAlignment": "Center"
# # #                     },
# # #                     {
# # #                         "type": "Column",
# # #                         "width": "stretch",
# # #                         "items": [
# # #                             {
# # #                                 "type": "Input.Text",
# # #                                 "id": f"training_text_{idx}",
# # #                                 "value": training,
# # #                                 "style": "text"
# # #                             }
# # #                         ],
# # #                         "verticalContentAlignment": "Center"
# # #                     },
# # #                     {
# # #                         "type": "Column",
# # #                         "width": "auto",
# # #                         "items": [
# # #                             {
# # #                                 "type": "ActionSet",
# # #                                 "actions": [
# # #                                     {
# # #                                         "type": "Action.Submit",
# # #                                         "title": "−",
# # #                                         "data": {
# # #                                             "action": "remove_training",
# # #                                             "id": idx
# # #                                         },
# # #                                         "style": "destructive"
# # #                                     }
# # #                                 ]
# # #                             }
# # #                         ],
# # #                         "verticalContentAlignment": "Center"
# # #                     }
# # #                 ],
# # #                 "spacing": "Small"
# # #             }
# # #             container["items"].append(training_row)
        
# # #         # Add Training Button
# # #         container["items"].append({
# # #             "type": "ActionSet",
# # #             "actions": [
# # #                 {
# # #                     "type": "Action.Submit",
# # #                     "title": "+ Add Training",
# # #                     "data": {
# # #                         "action": "add_training"
# # #                     }
# # #                 }
# # #             ],
# # #             "spacing": "Small"
# # #         })
        
# # #         # Additional Items Section
# # #         container["items"].append({
# # #             "type": "TextBlock",
# # #             "text": "Additional Items",
# # #             "weight": "Bolder",
# # #             "size": "Medium",
# # #             "spacing": "Large",
# # #             "wrap": True
# # #         })
        
# # #         # Additional items checkboxes with text inputs
# # #         for idx, item in enumerate(sections["additional_items"]):
# # #             item_row = {
# # #                 "type": "ColumnSet",
# # #                 "columns": [
# # #                     {
# # #                         "type": "Column",
# # #                         "width": "auto",
# # #                         "items": [
# # #                             {
# # #                                 "type": "Input.Toggle",
# # #                                 "id": f"additional_{idx}",
# # #                                 "title": "",
# # #                                 "value": "false"
# # #                             }
# # #                         ],
# # #                         "verticalContentAlignment": "Center"
# # #                     },
# # #                     {
# # #                         "type": "Column",
# # #                         "width": "stretch",
# # #                         "items": [
# # #                             {
# # #                                 "type": "Input.Text",
# # #                                 "id": f"additional_text_{idx}",
# # #                                 "placeholder": item,
# # #                                 "style": "text"
# # #                             }
# # #                         ],
# # #                         "verticalContentAlignment": "Center"
# # #                     }
# # #                 ],
# # #                 "spacing": "Small"
# # #             }
# # #             container["items"].append(item_row)
        
# # #         # Add Question Button
# # #         container["items"].append({
# # #             "type": "ActionSet",
# # #             "actions": [
# # #                 {
# # #                     "type": "Action.Submit",
# # #                     "title": "+ Add Question",
# # #                     "data": {
# # #                         "action": "add_question"
# # #                     }
# # #                 }
# # #             ],
# # #             "spacing": "Small"
# # #         })
        
# # #         card["body"].append(container)
        
# # #         # Main Actions at the bottom
# # #         card["actions"] = [
# # #             {
# # #                 "type": "Action.Submit",
# # #                 "title": "Save Template",
# # #                 "data": {
# # #                     "action": "save_template"
# # #                 },
# # #                 "style": "positive"
# # #             },
# # #             {
# # #                 "type": "Action.Submit",
# # #                 "title": "Close",
# # #                 "data": {
# # #                     "action": "close"
# # #                 }
# # #             }
# # #         ]
        
# # #         return card


# # # # Initialize generator
# # # generator = OnboardingCardGenerator()


# # # @app.route('/health', methods=['GET'])
# # # def health_check():
# # #     """Health check endpoint"""
# # #     return jsonify({
# # #         "status": "healthy",
# # #         "message": "Onboarding Card Generator API is running"
# # #     }), 200


# # # @app.route('/api/generate-card', methods=['POST'])
# # # def generate_card():
# # #     """
# # #     API endpoint to generate adaptive card from document
    
# # #     Request Body:
# # #     {
# # #         "document_path": "path/to/document.docx",
# # #         "return_metadata": false  // optional, default false
# # #     }
    
# # #     Response (return_metadata=false):
# # #     {
# # #         "$schema": "http://adaptivecards.io/schemas/adaptive-card.json",
# # #         "type": "AdaptiveCard",
# # #         ...
# # #     }
    
# # #     Response (return_metadata=true):
# # #     {
# # #         "success": true,
# # #         "adaptive_card": {...},
# # #         "metadata": {...}
# # #     }
# # #     """
# # #     try:
# # #         data = request.get_json()
        
# # #         if not data or 'document_path' not in data:
# # #             return jsonify({
# # #                 "error": "Missing 'document_path' in request body"
# # #             }), 400
        
# # #         document_path = data['document_path']
# # #         return_metadata = data.get('return_metadata', False)
        
# # #         # Validate file exists
# # #         if not os.path.exists(document_path):
# # #             return jsonify({
# # #                 "error": f"Document not found at path: {document_path}"
# # #             }), 404
        
# # #         # Read document
# # #         logger.info(f"Reading document from: {document_path}")
# # #         doc_content = generator.read_docx_file(document_path)
        
# # #         # Extract sections
# # #         sections = generator.extract_questions_from_document(doc_content)
        
# # #         # Generate adaptive card
# # #         adaptive_card = generator.generate_adaptive_card(sections)
        
# # #         # Return based on metadata flag
# # #         if return_metadata:
# # #             return jsonify({
# # #                 "success": True,
# # #                 "adaptive_card": adaptive_card,
# # #                 "metadata": {
# # #                     "document_path": document_path,
# # #                     "sections_count": {
# # #                         "general_info": len(sections["general_info"]),
# # #                         "trainings": len(sections["trainings"]),
# # #                         "additional_items": len(sections["additional_items"])
# # #                     }
# # #                 }
# # #             }), 200
# # #         else:
# # #             # Return clean adaptive card JSON directly
# # #             return jsonify(adaptive_card), 200
        
# # #     except FileNotFoundError as e:
# # #         logger.error(f"File not found: {str(e)}")
# # #         return jsonify({
# # #             "error": str(e)
# # #         }), 404
    
# # #     except Exception as e:
# # #         logger.error(f"Error generating card: {str(e)}")
# # #         return jsonify({
# # #             "error": str(e)
# # #         }), 500


# # # @app.route('/api/generate-card-from-sections', methods=['POST'])
# # # def generate_card_from_sections():
# # #     """
# # #     API endpoint to generate adaptive card from provided sections
    
# # #     Request Body:
# # #     {
# # #         "general_info": ["question 1", "question 2"],
# # #         "trainings": ["training 1", "training 2"],
# # #         "additional_items": ["item 1", "item 2"]
# # #     }
# # #     """
# # #     try:
# # #         data = request.get_json()
        
# # #         if not data:
# # #             return jsonify({
# # #                 "success": False,
# # #                 "error": "No data provided"
# # #             }), 400
        
# # #         sections = {
# # #             "general_info": data.get("general_info", []),
# # #             "trainings": data.get("trainings", []),
# # #             "additional_items": data.get("additional_items", [])
# # #         }
        
# # #         # Generate adaptive card
# # #         adaptive_card = generator.generate_adaptive_card(sections)
        
# # #         return jsonify({
# # #             "success": True,
# # #             "adaptive_card": adaptive_card
# # #         }), 200
        
# # #     except Exception as e:
# # #         logger.error(f"Error generating card from sections: {str(e)}")
# # #         return jsonify({
# # #             "success": False,
# # #             "error": str(e)
# # #         }), 500


# # # if __name__ == '__main__':
# # #     print("=" * 60)
# # #     print("Onboarding Adaptive Card Generator API")
# # #     print("=" * 60)
# # #     print("\nAvailable Endpoints:")
# # #     print("  GET  /health - Health check")
# # #     print("  POST /api/generate-card - Generate card from document")
# # #     print("  POST /api/generate-card-from-sections - Generate card from sections")
# # #     print("\nStarting server on http://localhost:5000")
# # #     print("=" * 60)
# # #     app.run(debug=True, host='0.0.0.0', port=5000)

# # from flask import Flask, request, jsonify
# # from flask_cors import CORS
# # import json
# # import re
# # from typing import List, Dict, Any
# # import os
# # from docx import Document
# # import logging

# # # Configure logging
# # logging.basicConfig(level=logging.INFO)
# # logger = logging.getLogger(__name__)

# # app = Flask(__name__)
# # CORS(app)  # Enable CORS for all routes

# # class OnboardingCardGenerator:
# #     """
# #     Generator for Microsoft Teams Adaptive Cards matching the exact template structure
# #     """
    
# #     def __init__(self):
# #         self.card_template = {
# #             "$schema": "http://adaptivecards.io/schemas/adaptive-card.json",
# #             "type": "AdaptiveCard",
# #             "version": "1.4",
# #             "body": [],
# #             "actions": []
# #         }
    
# #     def read_docx_file(self, filepath: str) -> str:
# #         """
# #         Read content from DOCX file
# #         """
# #         try:
# #             if not os.path.exists(filepath):
# #                 raise FileNotFoundError(f"File not found: {filepath}")
            
# #             doc = Document(filepath)
# #             full_text = []
            
# #             # Extract text from paragraphs
# #             for para in doc.paragraphs:
# #                 full_text.append(para.text)
            
# #             # Extract text from tables
# #             for table in doc.tables:
# #                 for row in table.rows:
# #                     row_text = []
# #                     for cell in row.cells:
# #                         row_text.append(cell.text)
# #                     full_text.append(' | '.join(row_text))
            
# #             content = '\n'.join(full_text)
# #             logger.info(f"Successfully read document: {filepath}")
# #             return content
            
# #         except Exception as e:
# #             logger.error(f"Error reading DOCX file: {str(e)}")
# #             raise
    
# #     def extract_questions_from_document(self, doc_content: str) -> Dict[str, List[str]]:
# #         """
# #         Extract questions from the actual document content
# #         Parse the table structure and extract items
# #         """
# #         sections = {
# #             "general_info": [],
# #             "trainings": [],
# #             "additional_items": []
# #         }
        
# #         lines = doc_content.split('\n')
# #         current_section = None
# #         in_mandatory_training = False
# #         in_general_section = False
# #         in_conversations = False
        
# #         for line in lines:
# #             line = line.strip()
            
# #             # Detect sections
# #             if 'General information' in line or 'to-do' in line:
# #                 current_section = 'general'
# #                 in_general_section = True
# #                 in_mandatory_training = False
# #                 in_conversations = False
# #                 continue
# #             elif 'Mandatory Training' in line:
# #                 current_section = 'trainings'
# #                 in_mandatory_training = True
# #                 in_general_section = False
# #                 in_conversations = False
# #                 continue
# #             elif 'Individual' in line and 'Conversation' in line:
# #                 current_section = 'conversations'
# #                 in_mandatory_training = False
# #                 in_general_section = False
# #                 in_conversations = True
# #                 continue
            
# #             # Extract general information items
# #             if in_general_section and '|' in line:
# #                 parts = [p.strip() for p in line.split('|') if p.strip()]
# #                 if parts and len(parts[0]) > 20 and not any(x in parts[0].lower() for x in ['content', 'description', 'responsible', 'duration', 'due date']):
# #                     item = parts[0].replace('**', '').strip()
# #                     if item and len(item) > 10 and item not in sections["general_info"]:
# #                         sections["general_info"].append(item)
            
# #             # Extract ALL training items with checkboxes
# #             if in_mandatory_training and '☐' in line:
# #                 match = re.search(r'☐\s*(.+?)(?:\s*\||$)', line)
# #                 if match:
# #                     training_name = match.group(1).strip()
# #                     training_name = re.sub(r'\[|\]|\{\.mark\}', '', training_name).strip()
# #                     if training_name and len(training_name) > 2 and training_name not in sections["trainings"]:
# #                         sections["trainings"].append(training_name)
            
# #             # Extract conversation/additional items
# #             if in_conversations and '☐' in line:
# #                 match = re.search(r'☐\s*(.+?)(?:\s*\||$)', line)
# #                 if match:
# #                     item_name = match.group(1).strip()
# #                     item_name = re.sub(r'\[|\]|\{\.mark\}|\*\*', '', item_name).strip()
# #                     if item_name and len(item_name) > 5 and item_name not in sections["additional_items"]:
# #                         sections["additional_items"].append(item_name)
        
# #         # Use defaults matching your document if nothing extracted
# #         if not sections["general_info"]:
# #             sections["general_info"] = [
# #                 "Have you attended the welcome event and received the welcome box?",
# #                 "Was there an introduction to your manager, colleagues, and buddy?",
# #                 "Is the workplace fully equipped and ready for work?"
# #             ]
        
# #         if not sections["trainings"]:
# #             sections["trainings"] = [
# #                 "Check-in @tesa",
# #                 "Anti-Corruption",
# #                 "Data-Protection",
# #                 "Cybersecurity",
# #                 "Code of Conduct",
# #                 "Safety and Health at Work",
# #                 "Sustainability Program@tesa",
# #                 "PME @tesa",
# #                 "tSN „to make The production of tesa tape",
# #                 "Other function related trainings"
# #             ]
        
# #         if not sections["additional_items"]:
# #             sections["additional_items"] = [
# #                 "Meeting with department stakeholders",
# #                 "Tour of facilities (offices, canteen, meeting areas)",
# #                 "Introduction to current projects and priorities"
# #             ]
        
# #         logger.info(f"Extracted sections: {len(sections['general_info'])} general, {len(sections['trainings'])} trainings, {len(sections['additional_items'])} additional")
# #         return sections
    
# #     def generate_adaptive_card(self, sections: Dict[str, List[str]]) -> Dict[str, Any]:
# #         """
# #         Generate the complete adaptive card JSON matching the exact template
# #         """
# #         card = {
# #             "$schema": "http://adaptivecards.io/schemas/adaptive-card.json",
# #             "type": "AdaptiveCard",
# #             "version": "1.4",
# #             "body": []
# #         }
        
# #         # Container with standard width
# #         container = {
# #             "type": "Container",
# #             "style": "default",
# #             "items": []
# #         }
        
# #         # Header Section
# #         header = {
# #             "type": "ColumnSet",
# #             "columns": [
# #                 {
# #                     "type": "Column",
# #                     "width": "auto",
# #                     "items": [
# #                         {
# #                             "type": "Image",
# #                             "url": "https://via.placeholder.com/50",
# #                             "size": "Small",
# #                             "style": "Person",
# #                             "altText": "NOA"
# #                         }
# #                     ],
# #                     "verticalContentAlignment": "Center"
# #                 },
# #                 {
# #                     "type": "Column",
# #                     "width": "stretch",
# #                     "items": [
# #                         {
# #                             "type": "TextBlock",
# #                             "text": "NOA",
# #                             "weight": "Bolder",
# #                             "size": "Medium",
# #                             "wrap": True
# #                         },
# #                         {
# #                             "type": "TextBlock",
# #                             "text": "Hiring Manager – Onboarding Template",
# #                             "isSubtle": True,
# #                             "spacing": "None",
# #                             "size": "Small",
# #                             "wrap": True
# #                         }
# #                     ],
# #                     "verticalContentAlignment": "Center"
# #                 }
# #             ],
# #             "separator": True,
# #             "spacing": "Medium"
# #         }
# #         container["items"].append(header)
        
# #         # General Information Section
# #         container["items"].append({
# #             "type": "TextBlock",
# #             "text": "General information / to-do's",
# #             "weight": "Bolder",
# #             "size": "Medium",
# #             "spacing": "Large",
# #             "wrap": True
# #         })
        
# #         # General checkboxes with editable text
# #         for idx, question in enumerate(sections["general_info"]):
# #             general_row = {
# #                 "type": "ColumnSet",
# #                 "columns": [
# #                     {
# #                         "type": "Column",
# #                         "width": "auto",
# #                         "items": [
# #                             {
# #                                 "type": "Input.Toggle",
# #                                 "id": f"general_{idx}",
# #                                 "title": "",
# #                                 "value": "false"
# #                             }
# #                         ],
# #                         "verticalContentAlignment": "Center"
# #                     },
# #                     {
# #                         "type": "Column",
# #                         "width": "stretch",
# #                         "items": [
# #                             {
# #                                 "type": "Input.Text",
# #                                 "id": f"general_text_{idx}",
# #                                 "value": question,
# #                                 "style": "text"
# #                             }
# #                         ],
# #                         "verticalContentAlignment": "Center"
# #                     }
# #                 ],
# #                 "spacing": "Small"
# #             }
# #             container["items"].append(general_row)
        
# #         # Trainings Section with nested checkboxes
# #         container["items"].append({
# #             "type": "Container",
# #             "items": [
# #                 {
# #                     "type": "Input.Toggle",
# #                     "id": "trainings_section",
# #                     "title": "Trainings",
# #                     "value": "false",
# #                     "wrap": True
# #                 }
# #             ],
# #             "spacing": "Medium"
# #         })
        
# #         # Training items with remove buttons
# #         for idx, training in enumerate(sections["trainings"]):
# #             training_row = {
# #                 "type": "ColumnSet",
# #                 "columns": [
# #                     {
# #                         "type": "Column",
# #                         "width": "auto",
# #                         "items": [
# #                             {
# #                                 "type": "Input.Toggle",
# #                                 "id": f"training_{idx}",
# #                                 "title": "",
# #                                 "value": "false"
# #                             }
# #                         ],
# #                         "verticalContentAlignment": "Center"
# #                     },
# #                     {
# #                         "type": "Column",
# #                         "width": "stretch",
# #                         "items": [
# #                             {
# #                                 "type": "Input.Text",
# #                                 "id": f"training_text_{idx}",
# #                                 "value": training,
# #                                 "style": "text"
# #                             }
# #                         ],
# #                         "verticalContentAlignment": "Center"
# #                     },
# #                     {
# #                         "type": "Column",
# #                         "width": "auto",
# #                         "items": [
# #                             {
# #                                 "type": "ActionSet",
# #                                 "actions": [
# #                                     {
# #                                         "type": "Action.Submit",
# #                                         "title": "−",
# #                                         "data": {
# #                                             "action": "remove_training",
# #                                             "id": idx
# #                                         },
# #                                         "style": "destructive"
# #                                     }
# #                                 ]
# #                             }
# #                         ],
# #                         "verticalContentAlignment": "Center"
# #                     }
# #                 ],
# #                 "spacing": "Small"
# #             }
# #             container["items"].append(training_row)
        
# #         # Add Training Button
# #         container["items"].append({
# #             "type": "ActionSet",
# #             "actions": [
# #                 {
# #                     "type": "Action.Submit",
# #                     "title": "+ Add Training",
# #                     "data": {
# #                         "action": "add_training"
# #                     }
# #                 }
# #             ],
# #             "spacing": "Small"
# #         })
        
# #         # Additional Items Section
# #         container["items"].append({
# #             "type": "TextBlock",
# #             "text": "Additional Items",
# #             "weight": "Bolder",
# #             "size": "Medium",
# #             "spacing": "Large",
# #             "wrap": True
# #         })
        
# #         # Additional items checkboxes with text inputs
# #         for idx, item in enumerate(sections["additional_items"]):
# #             item_row = {
# #                 "type": "ColumnSet",
# #                 "columns": [
# #                     {
# #                         "type": "Column",
# #                         "width": "auto",
# #                         "items": [
# #                             {
# #                                 "type": "Input.Toggle",
# #                                 "id": f"additional_{idx}",
# #                                 "title": "",
# #                                 "value": "false"
# #                             }
# #                         ],
# #                         "verticalContentAlignment": "Center"
# #                     },
# #                     {
# #                         "type": "Column",
# #                         "width": "stretch",
# #                         "items": [
# #                             {
# #                                 "type": "Input.Text",
# #                                 "id": f"additional_text_{idx}",
# #                                 "placeholder": item,
# #                                 "style": "text"
# #                             }
# #                         ],
# #                         "verticalContentAlignment": "Center"
# #                     }
# #                 ],
# #                 "spacing": "Small"
# #             }
# #             container["items"].append(item_row)
        
# #         # Add Question Button
# #         container["items"].append({
# #             "type": "ActionSet",
# #             "actions": [
# #                 {
# #                     "type": "Action.Submit",
# #                     "title": "+ Add Question",
# #                     "data": {
# #                         "action": "add_question"
# #                     }
# #                 }
# #             ],
# #             "spacing": "Small"
# #         })
        
# #         card["body"].append(container)
        
# #         # Main Actions at the bottom
# #         card["actions"] = [
# #             {
# #                 "type": "Action.Submit",
# #                 "title": "Save Template",
# #                 "data": {
# #                     "action": "save_template"
# #                 },
# #                 "style": "positive"
# #             },
# #             {
# #                 "type": "Action.Submit",
# #                 "title": "Close",
# #                 "data": {
# #                     "action": "close"
# #                 }
# #             }
# #         ]
        
# #         return card


# # # Initialize generator
# # generator = OnboardingCardGenerator()


# # @app.route('/health', methods=['GET'])
# # def health_check():
# #     """Health check endpoint"""
# #     return jsonify({
# #         "status": "healthy",
# #         "message": "Onboarding Card Generator API is running"
# #     }), 200


# # @app.route('/api/generate-card', methods=['POST'])
# # def generate_card():
# #     """
# #     API endpoint to generate adaptive card from document
    
# #     Request Body:
# #     {
# #         "document_path": "path/to/document.docx",
# #         "return_metadata": false  // optional, default false
# #     }
    
# #     Response (return_metadata=false):
# #     {
# #         "$schema": "http://adaptivecards.io/schemas/adaptive-card.json",
# #         "type": "AdaptiveCard",
# #         ...
# #     }
    
# #     Response (return_metadata=true):
# #     {
# #         "success": true,
# #         "adaptive_card": {...},
# #         "metadata": {...}
# #     }
# #     """
# #     try:
# #         data = request.get_json()
        
# #         if not data or 'document_path' not in data:
# #             return jsonify({
# #                 "error": "Missing 'document_path' in request body"
# #             }), 400
        
# #         document_path = data['document_path']
# #         return_metadata = data.get('return_metadata', False)
        
# #         # Validate file exists
# #         if not os.path.exists(document_path):
# #             return jsonify({
# #                 "error": f"Document not found at path: {document_path}"
# #             }), 404
        
# #         # Read document
# #         logger.info(f"Reading document from: {document_path}")
# #         doc_content = generator.read_docx_file(document_path)
        
# #         # Extract sections
# #         sections = generator.extract_questions_from_document(doc_content)
        
# #         # Generate adaptive card
# #         adaptive_card = generator.generate_adaptive_card(sections)
        
# #         # Return based on metadata flag
# #         if return_metadata:
# #             return jsonify({
# #                 "success": True,
# #                 "adaptive_card": adaptive_card,
# #                 "metadata": {
# #                     "document_path": document_path,
# #                     "sections_count": {
# #                         "general_info": len(sections["general_info"]),
# #                         "trainings": len(sections["trainings"]),
# #                         "additional_items": len(sections["additional_items"])
# #                     }
# #                 }
# #             }), 200
# #         else:
# #             # Return clean adaptive card JSON directly
# #             return jsonify(adaptive_card), 200
        
# #     except FileNotFoundError as e:
# #         logger.error(f"File not found: {str(e)}")
# #         return jsonify({
# #             "error": str(e)
# #         }), 404
    
# #     except Exception as e:
# #         logger.error(f"Error generating card: {str(e)}")
# #         return jsonify({
# #             "error": str(e)
# #         }), 500


# # @app.route('/api/generate-card-from-sections', methods=['POST'])
# # def generate_card_from_sections():
# #     """
# #     API endpoint to generate adaptive card from provided sections
    
# #     Request Body:
# #     {
# #         "general_info": ["question 1", "question 2"],
# #         "trainings": ["training 1", "training 2"],
# #         "additional_items": ["item 1", "item 2"]
# #     }
# #     """
# #     try:
# #         data = request.get_json()
        
# #         if not data:
# #             return jsonify({
# #                 "success": False,
# #                 "error": "No data provided"
# #             }), 400
        
# #         sections = {
# #             "general_info": data.get("general_info", []),
# #             "trainings": data.get("trainings", []),
# #             "additional_items": data.get("additional_items", [])
# #         }
        
# #         # Generate adaptive card
# #         adaptive_card = generator.generate_adaptive_card(sections)
        
# #         return jsonify({
# #             "success": True,
# #             "adaptive_card": adaptive_card
# #         }), 200
        
# #     except Exception as e:
# #         logger.error(f"Error generating card from sections: {str(e)}")
# #         return jsonify({
# #             "success": False,
# #             "error": str(e)
# #         }), 500


# # if __name__ == '__main__':
# #     print("=" * 60)
# #     print("Onboarding Adaptive Card Generator API")
# #     print("=" * 60)
# #     print("\nAvailable Endpoints:")
# #     print("  GET  /health - Health check")
# #     print("  POST /api/generate-card - Generate card from document")
# #     print("  POST /api/generate-card-from-sections - Generate card from sections")
# #     print("\nStarting server on http://localhost:5000")
# #     print("=" * 60)
# #     app.run(debug=True, host='0.0.0.0', port=5000)

# from flask import Flask, request, jsonify
# from flask_cors import CORS
# import json
# import re
# from typing import List, Dict, Any
# import os
# from docx import Document
# import logging

# # Configure logging
# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

# app = Flask(__name__)
# CORS(app)  # Enable CORS for all routes

# class OnboardingCardGenerator:
#     """
#     Generator for Microsoft Teams Adaptive Cards matching the exact template structure
#     """
    
#     def __init__(self):
#         self.card_template = {
#             "$schema": "https://adaptivecards.io/schemas/adaptive-card.json",
#             "type": "AdaptiveCard",
#             "version": "1.5",
#             "body": [],
#             "actions": []
#         }
    
#     def read_docx_file(self, filepath: str) -> str:
#         """
#         Read content from DOCX file
#         """
#         try:
#             if not os.path.exists(filepath):
#                 raise FileNotFoundError(f"File not found: {filepath}")
            
#             doc = Document(filepath)
#             full_text = []
            
#             # Extract text from paragraphs
#             for para in doc.paragraphs:
#                 full_text.append(para.text)
            
#             # Extract text from tables
#             for table in doc.tables:
#                 for row in table.rows:
#                     row_text = []
#                     for cell in row.cells:
#                         row_text.append(cell.text)
#                     full_text.append(' | '.join(row_text))
            
#             content = '\n'.join(full_text)
#             logger.info(f"Successfully read document: {filepath}")
#             return content
            
#         except Exception as e:
#             logger.error(f"Error reading DOCX file: {str(e)}")
#             raise
    
#     def extract_questions_from_document(self, doc_content: str) -> Dict[str, List[str]]:
#         """
#         Extract questions from the actual document content
#         Parse the table structure and extract items
#         """
#         sections = {
#             "general_info": [],
#             "trainings": [],
#             "additional_items": []
#         }
        
#         lines = doc_content.split('\n')
#         current_section = None
#         in_mandatory_training = False
#         in_general_section = False
#         in_conversations = False
        
#         for line in lines:
#             line = line.strip()
            
#             # Detect sections
#             if 'General information' in line or 'to-do' in line:
#                 current_section = 'general'
#                 in_general_section = True
#                 in_mandatory_training = False
#                 in_conversations = False
#                 continue
#             elif 'Mandatory Training' in line:
#                 current_section = 'trainings'
#                 in_mandatory_training = True
#                 in_general_section = False
#                 in_conversations = False
#                 continue
#             elif 'Individual' in line and 'Conversation' in line:
#                 current_section = 'conversations'
#                 in_mandatory_training = False
#                 in_general_section = False
#                 in_conversations = True
#                 continue
            
#             # Extract general information items
#             if in_general_section and '|' in line:
#                 parts = [p.strip() for p in line.split('|') if p.strip()]
#                 if parts and len(parts[0]) > 20 and not any(x in parts[0].lower() for x in ['content', 'description', 'responsible', 'duration', 'due date']):
#                     item = parts[0].replace('**', '').strip()
#                     if item and len(item) > 10 and item not in sections["general_info"]:
#                         sections["general_info"].append(item)
            
#             # Extract ALL training items with checkboxes
#             if in_mandatory_training and '☐' in line:
#                 match = re.search(r'☐\s*(.+?)(?:\s*\||$)', line)
#                 if match:
#                     training_name = match.group(1).strip()
#                     training_name = re.sub(r'\[|\]|\{\.mark\}', '', training_name).strip()
#                     if training_name and len(training_name) > 2 and training_name not in sections["trainings"]:
#                         sections["trainings"].append(training_name)
            
#             # Extract conversation/additional items
#             if in_conversations and '☐' in line:
#                 match = re.search(r'☐\s*(.+?)(?:\s*\||$)', line)
#                 if match:
#                     item_name = match.group(1).strip()
#                     item_name = re.sub(r'\[|\]|\{\.mark\}|\*\*', '', item_name).strip()
#                     if item_name and len(item_name) > 5 and item_name not in sections["additional_items"]:
#                         sections["additional_items"].append(item_name)
        
#         # Use defaults matching your document if nothing extracted
#         if not sections["general_info"]:
#             sections["general_info"] = [
#                 "Have you attended the welcome event and received the welcome box?",
#                 "Was there an introduction to your manager, colleagues, and buddy?",
#                 "Is the workplace fully equipped and ready for work?"
#             ]
        
#         if not sections["trainings"]:
#             sections["trainings"] = [
#                 "Check-in @tesa",
#                 "Anti-Corruption",
#                 "Data-Protection",
#                 "Cybersecurity",
#                 "Code of Conduct",
#                 "Safety and Health at Work",
#                 "Sustainability Program@tesa",
#                 "PME @tesa",
#                 "tSN „to make The production of tesa tape",
#                 "Other function related trainings"
#             ]
        
#         if not sections["additional_items"]:
#             sections["additional_items"] = [
#                 "Meeting with department stakeholders",
#                 "Tour of facilities (offices, canteen, meeting areas)",
#                 "Introduction to current projects and priorities"
#             ]
        
#         logger.info(f"Extracted sections: {len(sections['general_info'])} general, {len(sections['trainings'])} trainings, {len(sections['additional_items'])} additional")
#         return sections
    
#     def generate_adaptive_card(self, sections: Dict[str, List[str]]) -> Dict[str, Any]:
#         """
#         Generate the complete adaptive card JSON matching the exact template
#         """
#         card = {
#             "type": "AdaptiveCard",
#             "$schema": "https://adaptivecards.io/schemas/adaptive-card.json",
#             "version": "1.5",
#             "body": []
#         }
        
#         # Header Section
#         header = {
#             "type": "ColumnSet",
#             "columns": [
#                 {
#                     "type": "Column",
#                     "width": "auto",
#                     "items": [
#                         {
#                             "type": "Image",
#                             "url": "https://cdn-icons-png.flaticon.com/512/3135/3135715.png",
#                             "size": "Small",
#                             "altText": "Notification"
#                         }
#                     ]
#                 },
#                 {
#                     "type": "Column",
#                     "width": "stretch",
#                     "items": [
#                         {
#                             "type": "TextBlock",
#                             "text": "NOA",
#                             "weight": "Bolder",
#                             "size": "Medium",
#                             "color": "Accent"
#                         },
#                         {
#                             "type": "TextBlock",
#                             "text": "Hiring Manager – Onboarding Template",
#                             "isSubtle": True,
#                             "spacing": "None"
#                         }
#                     ]
#                 }
#             ],
#             "spacing": "Medium"
#         }
#         card["body"].append(header)
        
#         # General Information Section Title
#         card["body"].append({
#             "type": "TextBlock",
#             "text": "General information / to-do's",
#             "weight": "Bolder",
#             "size": "Large",
#             "spacing": "Large"
#         })
        
#         # General Information Container with emphasis style
#         general_container = {
#             "type": "Container",
#             "spacing": "Medium",
#             "style": "emphasis",
#             "items": []
#         }
        
#         # General checkboxes with editable text and checkbox images
#         for idx, question in enumerate(sections["general_info"]):
#             general_row = {
#                 "type": "ColumnSet",
#                 "columns": [
#                     {
#                         "type": "Column",
#                         "width": "auto",
#                         "spacing": "Small",
#                         "items": [
#                             {
#                                 "type": "Image",
#                                 "url": "data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='20' height='20'%3E%3Crect width='18' height='18' x='1' y='1' fill='%23e0e0e0' stroke='%23999' stroke-width='1'/%3E%3Cpath d='M5 10 L9 14 L15 6' stroke='%23666' stroke-width='2' fill='none'/%3E%3C/svg%3E",
#                                 "width": "20px",
#                                 "height": "20px"
#                             }
#                         ]
#                     },
#                     {
#                         "type": "Column",
#                         "width": "stretch",
#                         "items": [
#                             {
#                                 "type": "Input.Text",
#                                 "id": f"general_{idx + 1}_text",
#                                 "value": question
#                             }
#                         ]
#                     }
#                 ],
#                 "spacing": "Small" if idx > 0 else "None"
#             }
#             general_container["items"].append(general_row)
        
#         card["body"].append(general_container)
        
#         # Trainings Section - parent checkbox
#         trainings_header = {
#             "type": "ColumnSet",
#             "spacing": "Small",
#             "columns": [
#                 {
#                     "type": "Column",
#                     "width": "auto",
#                     "spacing": "Small",
#                     "items": [
#                         {
#                             "type": "Image",
#                             "url": "data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='20' height='20'%3E%3Crect width='18' height='18' x='1' y='1' fill='%23e0e0e0' stroke='%23999' stroke-width='1'/%3E%3Cpath d='M5 10 L9 14 L15 6' stroke='%23666' stroke-width='2' fill='none'/%3E%3C/svg%3E",
#                             "width": "20px",
#                             "height": "20px"
#                         }
#                     ]
#                 },
#                 {
#                     "type": "Column",
#                     "width": "stretch",
#                     "items": [
#                         {
#                             "type": "TextBlock",
#                             "text": "Trainings",
#                             "wrap": True,
#                             "weight": "Bolder"
#                         }
#                     ]
#                 }
#             ]
#         }
#         card["body"].append(trainings_header)
        
#         # Training items container (indented)
#         trainings_container = {
#             "type": "Container",
#             "spacing": "Small",
#             "style": "default",
#             "items": []
#         }
        
#         # Training items with indentation
#         for idx, training in enumerate(sections["trainings"]):
#             training_row = {
#                 "type": "ColumnSet",
#                 "columns": [
#                     {
#                         "type": "Column",
#                         "width": "20px",
#                         "items": [
#                             {
#                                 "type": "TextBlock",
#                                 "text": " ",
#                                 "wrap": True
#                             }
#                         ]
#                     },
#                     {
#                         "type": "Column",
#                         "width": "auto",
#                         "spacing": "Small",
#                         "items": [
#                             {
#                                 "type": "Input.Toggle",
#                                 "id": f"training_{idx + 1}"
#                             }
#                         ]
#                     },
#                     {
#                         "type": "Column",
#                         "width": "stretch",
#                         "items": [
#                             {
#                                 "type": "Input.Text",
#                                 "id": f"training_{idx + 1}_text",
#                                 "value": training
#                             }
#                         ]
#                     },
#                     {
#                         "type": "Column",
#                         "width": "auto",
#                         "items": [
#                             {
#                                 "type": "ActionSet",
#                                 "actions": [
#                                     {
#                                         "type": "Action.Submit",
#                                         "title": "−",
#                                         "data": {
#                                             "action": "removeTraining",
#                                             "id": f"training_{idx + 1}"
#                                         }
#                                     }
#                                 ]
#                             }
#                         ]
#                     }
#                 ],
#                 "spacing": "Small" if idx > 0 else "None"
#             }
#             trainings_container["items"].append(training_row)
        
#         # Add Training button (indented)
#         add_training_row = {
#             "type": "ColumnSet",
#             "spacing": "Small",
#             "columns": [
#                 {
#                     "type": "Column",
#                     "width": "20px",
#                     "items": [
#                         {
#                             "type": "TextBlock",
#                             "text": " ",
#                             "wrap": True
#                         }
#                     ]
#                 },
#                 {
#                     "type": "Column",
#                     "width": "stretch",
#                     "items": [
#                         {
#                             "type": "ActionSet",
#                             "actions": [
#                                 {
#                                     "type": "Action.Submit",
#                                     "title": "✚ Add Training",
#                                     "data": {
#                                         "action": "addTraining"
#                                     }
#                                 }
#                             ]
#                         }
#                     ]
#                 }
#             ]
#         }
#         trainings_container["items"].append(add_training_row)
        
#         card["body"].append(trainings_container)
        
#         # Additional Items Section
#         card["body"].append({
#             "type": "TextBlock",
#             "text": "Additional Items",
#             "weight": "Bolder",
#             "size": "Medium",
#             "spacing": "Large"
#         })
        
#         # Additional items container
#         additional_container = {
#             "type": "Container",
#             "spacing": "Small",
#             "items": []
#         }
        
#         # Additional items checkboxes with text inputs
#         for idx, item in enumerate(sections["additional_items"]):
#             item_row = {
#                 "type": "ColumnSet",
#                 "columns": [
#                     {
#                         "type": "Column",
#                         "width": "auto",
#                         "spacing": "Small",
#                         "items": [
#                             {
#                                 "type": "Input.Toggle",
#                                 "id": f"additional_{idx + 1}"
#                             }
#                         ]
#                     },
#                     {
#                         "type": "Column",
#                         "width": "stretch",
#                         "items": [
#                             {
#                                 "type": "Input.Text",
#                                 "id": f"additional_{idx + 1}_text",
#                                 "value": item
#                             }
#                         ]
#                     }
#                 ],
#                 "spacing": "Small" if idx > 0 else "None"
#             }
#             additional_container["items"].append(item_row)
        
#         card["body"].append(additional_container)
        
#         # Add Question Button
#         card["body"].append({
#             "type": "ActionSet",
#             "spacing": "Small",
#             "actions": [
#                 {
#                     "type": "Action.Submit",
#                     "title": "✚ Add Question",
#                     "data": {
#                         "action": "addQuestion"
#                     }
#                 }
#             ]
#         })
        
#         # Main Actions at the bottom
#         card["actions"] = [
#             {
#                 "type": "Action.Submit",
#                 "title": "💾 Save Template",
#                 "data": {
#                     "action": "save"
#                 },
#                 "style": "positive"
#             },
#             {
#                 "type": "Action.Submit",
#                 "title": "Close",
#                 "data": {
#                     "action": "close"
#                 }
#             }
#         ]
        
#         return card


# # Initialize generator
# generator = OnboardingCardGenerator()


# @app.route('/health', methods=['GET'])
# def health_check():
#     """Health check endpoint"""
#     return jsonify({
#         "status": "healthy",
#         "message": "Onboarding Card Generator API is running"
#     }), 200


# @app.route('/api/generate-card', methods=['POST'])
# def generate_card():
#     """
#     API endpoint to generate adaptive card from document
    
#     Request Body:
#     {
#         "document_path": "path/to/document.docx",
#         "return_metadata": false  // optional, default false
#     }
    
#     Response (return_metadata=false):
#     {
#         "$schema": "https://adaptivecards.io/schemas/adaptive-card.json",
#         "type": "AdaptiveCard",
#         ...
#     }
    
#     Response (return_metadata=true):
#     {
#         "success": true,
#         "adaptive_card": {...},
#         "metadata": {...}
#     }
#     """
#     try:
#         data = request.get_json()
        
#         if not data or 'document_path' not in data:
#             return jsonify({
#                 "error": "Missing 'document_path' in request body"
#             }), 400
        
#         document_path = data['document_path']
#         return_metadata = data.get('return_metadata', False)
        
#         # Validate file exists
#         if not os.path.exists(document_path):
#             return jsonify({
#                 "error": f"Document not found at path: {document_path}"
#             }), 404
        
#         # Read document
#         logger.info(f"Reading document from: {document_path}")
#         doc_content = generator.read_docx_file(document_path)
        
#         # Extract sections
#         sections = generator.extract_questions_from_document(doc_content)
        
#         # Generate adaptive card
#         adaptive_card = generator.generate_adaptive_card(sections)
        
#         # Return based on metadata flag
#         if return_metadata:
#             return jsonify({
#                 "success": True,
#                 "adaptive_card": adaptive_card,
#                 "metadata": {
#                     "document_path": document_path,
#                     "sections_count": {
#                         "general_info": len(sections["general_info"]),
#                         "trainings": len(sections["trainings"]),
#                         "additional_items": len(sections["additional_items"])
#                     }
#                 }
#             }), 200
#         else:
#             # Return clean adaptive card JSON directly
#             return jsonify(adaptive_card), 200
        
#     except FileNotFoundError as e:
#         logger.error(f"File not found: {str(e)}")
#         return jsonify({
#             "error": str(e)
#         }), 404
    
#     except Exception as e:
#         logger.error(f"Error generating card: {str(e)}")
#         return jsonify({
#             "error": str(e)
#         }), 500


# @app.route('/api/generate-card-from-sections', methods=['POST'])
# def generate_card_from_sections():
#     """
#     API endpoint to generate adaptive card from provided sections
    
#     Request Body:
#     {
#         "general_info": ["question 1", "question 2"],
#         "trainings": ["training 1", "training 2"],
#         "additional_items": ["item 1", "item 2"]
#     }
#     """
#     try:
#         data = request.get_json()
        
#         if not data:
#             return jsonify({
#                 "error": "No data provided"
#             }), 400
        
#         sections = {
#             "general_info": data.get("general_info", []),
#             "trainings": data.get("trainings", []),
#             "additional_items": data.get("additional_items", [])
#         }
        
#         # Generate adaptive card
#         adaptive_card = generator.generate_adaptive_card(sections)
        
#         return jsonify(adaptive_card), 200
        
#     except Exception as e:
#         logger.error(f"Error generating card from sections: {str(e)}")
#         return jsonify({
#             "error": str(e)
#         }), 500


# if __name__ == '__main__':
#     print("=" * 60)
#     print("Onboarding Adaptive Card Generator API")
#     print("=" * 60)
#     print("\nAvailable Endpoints:")
#     print("  GET  /health - Health check")
#     print("  POST /api/generate-card - Generate card from document")
#     print("  POST /api/generate-card-from-sections - Generate card from sections")
#     print("\nStarting server on http://localhost:5000")
#     print("=" * 60)
#     app.run(debug=True, host='0.0.0.0', port=5000)

from flask import Flask, request, jsonify
from flask_cors import CORS
import json
import re
from typing import List, Dict, Any
import os
from docx import Document
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

class OnboardingCardGenerator:
    """
    Generator for Microsoft Teams Adaptive Cards - Fully Dynamic
    Extracts sections and content dynamically from any document structure
    """
    
    def __init__(self):
        self.card_template = {
            "$schema": "https://adaptivecards.io/schemas/adaptive-card.json",
            "type": "AdaptiveCard",
            "version": "1.5",
            "body": [],
            "actions": []
        }
    
    def read_docx_file(self, filepath: str) -> str:
        """
        Read content from DOCX file
        """
        try:
            if not os.path.exists(filepath):
                raise FileNotFoundError(f"File not found: {filepath}")
            
            doc = Document(filepath)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    full_text.append(' | '.join(row_text))
            
            content = '\n'.join(full_text)
            logger.info(f"Successfully read document: {filepath}")
            return content
            
        except Exception as e:
            logger.error(f"Error reading DOCX file: {str(e)}")
            raise
    
    def extract_sections_dynamically(self, doc_content: str) -> Dict[str, Any]:
        """
        Dynamically extract ALL sections and their items from the document
        Returns a dictionary with section names as keys and lists of items as values
        """
        sections = {}
        lines = doc_content.split('\n')
        current_section = None
        current_section_key = None
        section_counter = 0
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            # Skip empty lines, table borders, and headers
            if not line or line.startswith('+') or line.startswith('='):
                continue
            
            # Detect section headers - look for bold text or specific patterns
            # Section headers typically contain: **, specific keywords, or are in all caps
            is_section_header = False
            section_name = None
            
            # Check for various section header patterns
            if '**' in line and ('information' in line.lower() or 'training' in line.lower() or 
                                 'conversation' in line.lower() or 'item' in line.lower() or
                                 'to-do' in line.lower()):
                # Extract section name from markdown bold
                section_name = re.sub(r'\*\*|\||{\.mark}', '', line).strip()
                is_section_header = True
            elif line.isupper() and len(line) > 5:
                # All caps might indicate a header
                section_name = line
                is_section_header = True
            elif any(keyword in line.lower() for keyword in [
                'general information', 'mandatory training', 'individual conversation',
                'additional items', 'trainings', 'to-do'
            ]):
                section_name = re.sub(r'\*\*|\||{\.mark}', '', line).strip()
                is_section_header = True
            
            # If we found a section header, create a new section
            if is_section_header and section_name:
                section_counter += 1
                # Create a clean key for the section
                current_section_key = f"section_{section_counter}"
                current_section = section_name
                sections[current_section_key] = {
                    "name": section_name,
                    "items": [],
                    "type": self._determine_section_type(section_name)
                }
                logger.info(f"Found section: '{section_name}' (key: {current_section_key})")
                continue
            
            # Extract items from the current section
            if current_section_key:
                item = None
                
                # Look for checkbox items (☐)
                if '☐' in line:
                    match = re.search(r'☐\s*(.+?)(?:\s*\||$)', line)
                    if match:
                        item = match.group(1).strip()
                        item = re.sub(r'\[|\]|\{\.mark\}|\*\*', '', item).strip()
                
                # Look for table row items
                elif '|' in line and not line.startswith('|'):
                    parts = [p.strip() for p in line.split('|') if p.strip()]
                    if parts:
                        # Get the first meaningful column (usually the description)
                        potential_item = parts[0].replace('**', '').replace('*', '').strip()
                        # Filter out headers and metadata
                        if (len(potential_item) > 10 and 
                            not any(x in potential_item.lower() for x in [
                                'content', 'description', 'responsible', 'duration', 
                                'due date', 'days', 'respo'
                            ])):
                            item = potential_item
                
                # Add the item if we found one
                if item and len(item) > 5:
                    # Avoid duplicates
                    if item not in sections[current_section_key]["items"]:
                        sections[current_section_key]["items"].append(item)
                        logger.info(f"  Added item to {current_section_key}: {item[:50]}...")
        
        # Log extraction summary
        logger.info(f"\nExtraction Summary:")
        for key, section in sections.items():
            logger.info(f"  {key} ({section['name']}): {len(section['items'])} items")
        
        return sections
    
    def _determine_section_type(self, section_name: str) -> str:
        """
        Determine the type of section based on its name
        Returns: 'general', 'trainings', 'additional', or 'custom'
        """
        section_lower = section_name.lower()
        
        if any(keyword in section_lower for keyword in ['general', 'information', 'to-do']):
            return 'general'
        elif any(keyword in section_lower for keyword in ['training', 'mandatory']):
            return 'trainings'
        elif any(keyword in section_lower for keyword in ['additional', 'conversation', 'individual']):
            return 'additional'
        else:
            return 'custom'
    
    def generate_adaptive_card(self, sections: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate adaptive card dynamically from extracted sections
        """
        card = {
            "type": "AdaptiveCard",
            "$schema": "https://adaptivecards.io/schemas/adaptive-card.json",
            "version": "1.5",
            "body": []
        }
        
        # Header Section - Always present
        header = {
            "type": "ColumnSet",
            "columns": [
                {
                    "type": "Column",
                    "width": "auto",
                    "items": [
                        {
                            "type": "Image",
                            "url": "https://cdn-icons-png.flaticon.com/512/3135/3135715.png",
                            "size": "Small",
                            "altText": "Notification"
                        }
                    ]
                },
                {
                    "type": "Column",
                    "width": "stretch",
                    "items": [
                        {
                            "type": "TextBlock",
                            "text": "NOA",
                            "weight": "Bolder",
                            "size": "Medium",
                            "color": "Accent"
                        },
                        {
                            "type": "TextBlock",
                            "text": "Hiring Manager – Onboarding Template",
                            "isSubtle": True,
                            "spacing": "None"
                        }
                    ]
                }
            ],
            "spacing": "Medium"
        }
        card["body"].append(header)
        
        # Process each section dynamically
        for section_key, section_data in sections.items():
            section_name = section_data["name"]
            section_items = section_data["items"]
            section_type = section_data["type"]
            
            if not section_items:
                continue
            
            # Add section based on its type
            if section_type == 'general':
                self._add_general_section(card["body"], section_name, section_items, section_key)
            elif section_type == 'trainings':
                self._add_trainings_section(card["body"], section_name, section_items, section_key)
            elif section_type == 'additional':
                self._add_additional_section(card["body"], section_name, section_items, section_key)
            else:
                # Custom section - use additional items style
                self._add_custom_section(card["body"], section_name, section_items, section_key)
        
        # Main Actions - Always present
        card["actions"] = [
            {
                "type": "Action.Submit",
                "title": "💾 Save Template",
                "data": {
                    "action": "save"
                },
                "style": "positive"
            },
            {
                "type": "Action.Submit",
                "title": "Close",
                "data": {
                    "action": "close"
                }
            }
        ]
        
        return card
    
    def _add_general_section(self, body: List, section_name: str, items: List[str], section_key: str):
        """Add a general information section with emphasis style"""
        body.append({
            "type": "TextBlock",
            "text": section_name,
            "weight": "Bolder",
            "size": "Large",
            "spacing": "Large"
        })
        
        general_container = {
            "type": "Container",
            "spacing": "Medium",
            "style": "emphasis",
            "items": []
        }
        
        for idx, question in enumerate(items):
            general_row = {
                "type": "ColumnSet",
                "columns": [
                    {
                        "type": "Column",
                        "width": "auto",
                        "spacing": "Small",
                        "items": [
                            {
                                "type": "Image",
                                "url": "data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='20' height='20'%3E%3Crect width='18' height='18' x='1' y='1' fill='%23e0e0e0' stroke='%23999' stroke-width='1'/%3E%3Cpath d='M5 10 L9 14 L15 6' stroke='%23666' stroke-width='2' fill='none'/%3E%3C/svg%3E",
                                "width": "20px",
                                "height": "20px"
                            }
                        ]
                    },
                    {
                        "type": "Column",
                        "width": "stretch",
                        "items": [
                            {
                                "type": "Input.Text",
                                "id": f"{section_key}_{idx + 1}_text",
                                "value": question
                            }
                        ]
                    }
                ],
                "spacing": "Small" if idx > 0 else "None"
            }
            general_container["items"].append(general_row)
        
        body.append(general_container)
    
    def _add_trainings_section(self, body: List, section_name: str, items: List[str], section_key: str):
        """Add a trainings section with parent checkbox and indented items"""
        trainings_header = {
            "type": "ColumnSet",
            "spacing": "Small",
            "columns": [
                {
                    "type": "Column",
                    "width": "auto",
                    "spacing": "Small",
                    "items": [
                        {
                            "type": "Image",
                            "url": "data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='20' height='20'%3E%3Crect width='18' height='18' x='1' y='1' fill='%23e0e0e0' stroke='%23999' stroke-width='1'/%3E%3Cpath d='M5 10 L9 14 L15 6' stroke='%23666' stroke-width='2' fill='none'/%3E%3C/svg%3E",
                            "width": "20px",
                            "height": "20px"
                        }
                    ]
                },
                {
                    "type": "Column",
                    "width": "stretch",
                    "items": [
                        {
                            "type": "TextBlock",
                            "text": section_name,
                            "wrap": True,
                            "weight": "Bolder"
                        }
                    ]
                }
            ]
        }
        body.append(trainings_header)
        
        trainings_container = {
            "type": "Container",
            "spacing": "Small",
            "style": "default",
            "items": []
        }
        
        for idx, training in enumerate(items):
            training_row = {
                "type": "ColumnSet",
                "columns": [
                    {
                        "type": "Column",
                        "width": "20px",
                        "items": [
                            {
                                "type": "TextBlock",
                                "text": " ",
                                "wrap": True
                            }
                        ]
                    },
                    {
                        "type": "Column",
                        "width": "auto",
                        "spacing": "Small",
                        "items": [
                            {
                                "type": "Input.Toggle",
                                "id": f"{section_key}_{idx + 1}"
                            }
                        ]
                    },
                    {
                        "type": "Column",
                        "width": "stretch",
                        "items": [
                            {
                                "type": "Input.Text",
                                "id": f"{section_key}_{idx + 1}_text",
                                "value": training
                            }
                        ]
                    },
                    {
                        "type": "Column",
                        "width": "auto",
                        "items": [
                            {
                                "type": "ActionSet",
                                "actions": [
                                    {
                                        "type": "Action.Submit",
                                        "title": "−",
                                        "data": {
                                            "action": "removeTraining",
                                            "id": f"{section_key}_{idx + 1}"
                                        }
                                    }
                                ]
                            }
                        ]
                    }
                ],
                "spacing": "Small" if idx > 0 else "None"
            }
            trainings_container["items"].append(training_row)
        
        # Add Training button
        add_training_row = {
            "type": "ColumnSet",
            "spacing": "Small",
            "columns": [
                {
                    "type": "Column",
                    "width": "20px",
                    "items": [
                        {
                            "type": "TextBlock",
                            "text": " ",
                            "wrap": True
                        }
                    ]
                },
                {
                    "type": "Column",
                    "width": "stretch",
                    "items": [
                        {
                            "type": "ActionSet",
                            "actions": [
                                {
                                    "type": "Action.Submit",
                                    "title": "✚ Add Training",
                                    "data": {
                                        "action": "addTraining",
                                        "section": section_key
                                    }
                                }
                            ]
                        }
                    ]
                }
            ]
        }
        trainings_container["items"].append(add_training_row)
        body.append(trainings_container)
    
    def _add_additional_section(self, body: List, section_name: str, items: List[str], section_key: str):
        """Add an additional items section"""
        body.append({
            "type": "TextBlock",
            "text": section_name,
            "weight": "Bolder",
            "size": "Medium",
            "spacing": "Large"
        })
        
        additional_container = {
            "type": "Container",
            "spacing": "Small",
            "items": []
        }
        
        for idx, item in enumerate(items):
            item_row = {
                "type": "ColumnSet",
                "columns": [
                    {
                        "type": "Column",
                        "width": "auto",
                        "spacing": "Small",
                        "items": [
                            {
                                "type": "Input.Toggle",
                                "id": f"{section_key}_{idx + 1}"
                            }
                        ]
                    },
                    {
                        "type": "Column",
                        "width": "stretch",
                        "items": [
                            {
                                "type": "Input.Text",
                                "id": f"{section_key}_{idx + 1}_text",
                                "value": item
                            }
                        ]
                    }
                ],
                "spacing": "Small" if idx > 0 else "None"
            }
            additional_container["items"].append(item_row)
        
        body.append(additional_container)
        
        # Add Question Button
        body.append({
            "type": "ActionSet",
            "spacing": "Small",
            "actions": [
                {
                    "type": "Action.Submit",
                    "title": "✚ Add Question",
                    "data": {
                        "action": "addQuestion",
                        "section": section_key
                    }
                }
            ]
        })
    
    def _add_custom_section(self, body: List, section_name: str, items: List[str], section_key: str):
        """Add a custom section (uses additional items style)"""
        self._add_additional_section(body, section_name, items, section_key)


# Initialize generator
generator = OnboardingCardGenerator()


@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        "status": "healthy",
        "message": "Onboarding Card Generator API is running"
    }), 200


@app.route('/api/generate-card', methods=['POST'])
def generate_card():
    """
    API endpoint to generate adaptive card from document
    
    Request Body:
    {
        "document_path": "path/to/document.docx",
        "return_metadata": false  // optional, default false
    }
    """
    try:
        data = request.get_json()
        
        if not data or 'document_path' not in data:
            return jsonify({
                "error": "Missing 'document_path' in request body"
            }), 400
        
        document_path = data['document_path']
        return_metadata = data.get('return_metadata', False)
        
        # Validate file exists
        if not os.path.exists(document_path):
            return jsonify({
                "error": f"Document not found at path: {document_path}"
            }), 404
        
        # Read document
        logger.info(f"Reading document from: {document_path}")
        doc_content = generator.read_docx_file(document_path)
        
        # Extract sections dynamically
        sections = generator.extract_sections_dynamically(doc_content)
        
        # Generate adaptive card
        adaptive_card = generator.generate_adaptive_card(sections)
        
        # Return based on metadata flag
        if return_metadata:
            metadata = {
                "document_path": document_path,
                "sections": {}
            }
            for key, section in sections.items():
                metadata["sections"][key] = {
                    "name": section["name"],
                    "type": section["type"],
                    "item_count": len(section["items"])
                }
            
            return jsonify({
                "success": True,
                "adaptive_card": adaptive_card,
                "metadata": metadata
            }), 200
        else:
            # Return clean adaptive card JSON directly
            return jsonify(adaptive_card), 200
        
    except FileNotFoundError as e:
        logger.error(f"File not found: {str(e)}")
        return jsonify({
            "error": str(e)
        }), 404
    
    except Exception as e:
        logger.error(f"Error generating card: {str(e)}")
        return jsonify({
            "error": str(e)
        }), 500


@app.route('/api/generate-card-from-sections', methods=['POST'])
def generate_card_from_sections():
    """
    API endpoint to generate adaptive card from provided sections
    
    Request Body:
    {
        "sections": {
            "section_1": {
                "name": "General Information",
                "type": "general",
                "items": ["item1", "item2"]
            },
            "section_2": {
                "name": "Trainings",
                "type": "trainings",
                "items": ["training1", "training2"]
            }
        }
    }
    """
    try:
        data = request.get_json()
        
        if not data or 'sections' not in data:
            return jsonify({
                "error": "Missing 'sections' in request body"
            }), 400
        
        sections = data['sections']
        
        # Generate adaptive card
        adaptive_card = generator.generate_adaptive_card(sections)
        
        return jsonify(adaptive_card), 200
        
    except Exception as e:
        logger.error(f"Error generating card from sections: {str(e)}")
        return jsonify({
            "error": str(e)
        }), 500


if __name__ == '__main__':
    print("=" * 60)
    print("Onboarding Adaptive Card Generator API - Fully Dynamic")
    print("=" * 60)
    print("\nAvailable Endpoints:")
    print("  GET  /health - Health check")
    print("  POST /api/generate-card - Generate card from document")
    print("  POST /api/generate-card-from-sections - Generate card from sections")
    print("\nFeatures:")
    print("  ✓ Dynamically extracts ALL sections from document")
    print("  ✓ No hardcoded section names or defaults")
    print("  ✓ Automatically determines section types")
    print("  ✓ Handles any document structure")
    print("\nStarting server on http://localhost:5000")
    print("=" * 60)
    app.run(debug=True, host='0.0.0.0', port=5000)